import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import ollama
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pdfplumber
import json
from ollama import ResponseError
import re
from json_repair import repair_json
import random
META_PHRASES = [
    "here is",
    "here's",
    "rewritten resume summary",
    "based on the provided text",
    "below is",
    "following is"
]

def remove_meta_text(text: str) -> str:
    lines = text.splitlines()
    clean_lines = []

    for line in lines:
        lower = line.lower()
        if any(p in lower for p in META_PHRASES):
            continue
        clean_lines.append(line)

    return " ".join(clean_lines).strip()
VARIATION_STYLES = [
    "professional and concise",
    "calm and neutral",
    "confident but simple",
    "reflective and academic",
    "straightforward and ATS-friendly"
]

FORBIDDEN_TERMS = {
    "results-driven", "business growth", "high-pressure",
    "stakeholders", "driving success", "competitive edge",
    "industry", "organization", "company", "leader", "outstanding",
    "expert", "business growth"
}
SOFT_SKILL_KEYWORDS = {
    "communication", "teamwork", "leadership", "problem solving",
    "time management", "adaptability", "critical thinking",
    "creativity", "collaboration", "work ethic", "flexibility",
    "decision making", "emotional intelligence", "interpersonal",
    "project management",
    "public relations"
}

POSITIVE_TRAITS = {
    "brilliant", "smart", "hardworking", "dedicated", "motivated",
    "passionate", "focused", "quick", "learner", "creative",
    "disciplined", "confident", "adaptable", "responsible"
}

JUNK_WORDS = {
    "ok", "okay", "good", "fine", "nice", "great",
    "yes", "no", "cool", "awesome", "nothing","i am good"
}


def is_intent_based_summary(text: str) -> bool:
    if not text or not text.strip():
        return False

    text = text.lower().strip()
    words = set(re.findall(r"[a-z]+", text))

    # Must indicate self-description
    has_self_reference = any(p in text for p in ["i am", "i'm", "iam"])

    # Must contain positive intent
    has_positive_trait = bool(words & POSITIVE_TRAITS)

    return has_self_reference and has_positive_trait
BANNED_WORDS = {
    "skilled", "experience", "experienced", "expert", "expertise",
    "proven", "ability", "abilities", "capable", "talented",
    "versatile", "dedicated", "strong", "excellent",
    "strategic", "strategy", "team", "coordination",
    "competitive", "performance", "adapt", "excel"
}
def classify_input(text: str) -> str:
    text = text.lower()

    experience_keywords = ["years", "worked", "experience", "responsible for"]
    skill_keywords = ["skill", "knowledge of", "trained in", "proficient in"]
    interest_keywords = ["like", "enjoy", "interest", "hobby", "good in"]

    if any(k in text for k in experience_keywords):
        return "experience"
    if any(k in text for k in skill_keywords):
        return "skill"
    if any(k in text for k in interest_keywords):
        return "interest"
    return "neutral"

def sanitize_summary(text: str) -> str:
    for term in FORBIDDEN_TERMS:
        text = re.sub(rf"\b{term}\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()
def is_invalid_summary(text: str) -> bool:
    # reject very short or broken output
    if len(text.split()) < 12:
        return True
    if not text[0].isupper():
        return True
    if "." not in text:
        return True
    return False

def is_low_quality_summary(text: str) -> bool:
    if not text or not text.strip():
        return True

    text = text.lower().strip()
    words = re.findall(r"[a-z]+", text)

    if len(text) < 20:
        return True

    if len(words) < 4:
        return True

    if all(word in JUNK_WORDS for word in words):
        return True

    if len(set(words)) <= 2:
        return True

    return False

def extract_contact_regex(text):
    email = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    phone = re.findall(r"\b\d{10}\b", text)

    name = ""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        name = lines[0]  # heuristic

    return {
        "name": name,
        "email": email[0] if email else "",
        "phone": phone[0] if phone else ""
    }
def extract_location_safely(resume_text):
    lines = resume_text.splitlines()

    # only top 15 lines
    header_text = " ".join(lines[:15])

    # simple city-state-country patterns
    patterns = [
        r"[A-Z][a-z]+,\s*[A-Z]{2}",          # Los Angeles, CA
        r"[A-Z][a-z]+\s+[A-Z][a-z]+,\s*[A-Z]{2}",
        r"[A-Z][a-z]+,\s*[A-Z][a-z]+"        # San Antonio, Texas
    ]

    for p in patterns:
        match = re.search(p, header_text)
        if match:
            return match.group()

    return ""
# functions
def generate_ai_content(prompt):
    try:
        response = ollama.chat(model="llama3.2:latest", messages=[
            {"role": "system", "content": "You are an expert resume writer. Provide ONLY the requested content. No conversational filler like 'Here is your summary'."},
            {"role": "user", "content": prompt}
        ])
        return response['message']['content'].strip()
    except Exception as e:
        return f"Error: {str(e)}"
#used for pdf or word
def extract_resume_text(file):
    text = ""

    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"

    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)

        # paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # tables (SIDEBAR FIX)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    return text
#autofill
def ats_parse_resume(resume_text):
    resume_text = resume_text[:6000]

    prompt = f"""
Return ONLY valid JSON.

Schema:
{{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "summary": "",

  "education": [
    {{
      "course": "",
      "school": "",
      "board": "",
      "startyear": "",
      "stopyear": "",
      "sgpa": ""
    }}
  ],

  "skills_list": [],
  "languages": [],
  "soft_options": [],

  "experience_raw": "",
  "projects_raw": "",
  "declaration_raw": ""
}}

Resume:
\"\"\"{resume_text}\"\"\"
"""

    response = ollama.chat(
        model="llama3.2:latest",
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response["message"]["content"]
    #clean = raw[raw.find("{"): raw.rfind("}")+1]
    fixed_json = repair_json(raw)

    try:
        result = json.loads(fixed_json)
    except Exception:
        result = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "",
        "education": [],
        "skills_list": [],
        "languages": [],
        "soft_options": [],
        "experience_raw": "",
        "projects_raw": "",
        "declaration_raw": ""
    }

    return result
#used for mismatch structure
def normalize_ats_data(p):
    if isinstance(p, list):
        p = {
            "name": "",
            "email": "",
            "phone": "",
            "location": "",
            "summary": "",
            "education": [],
            "skills_list": [],
            "languages": [],
            "soft_options": [],
            "experience_raw": "",
            "projects_raw": "",
            "declaration_raw": ""
        }

    # 🔑 FIX 2: ATS returned None or garbage
    if not isinstance(p, dict):
        p = {}

    # 🔧 PROFILE → SUMMARY FALLBACK
    if not p.get("summary") and p.get("profile"):
        if isinstance(p["profile"], str):
            p["summary"] = p["profile"]
        elif isinstance(p["profile"], dict):
            p["summary"] = p["profile"].get("text", "")

    # 🔧 FIX SUMMARY
    summary = p.get("summary", "")
    if isinstance(summary, dict):
        p["summary"] = summary.get("text", "")
    elif isinstance(summary, list):
        p["summary"] = " ".join(map(str, summary))
    else:
        p["summary"] = str(summary)

    # 🔑 FIX 3: Ensure all keys exist
    defaults = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        #"summary": "",
        "education": [],
        "skills_list": [],
        "languages": [],
        "soft_options": [],
        "experience_raw": "",
        "projects_raw": "",
        "declaration_raw": ""
     }

    # Lists
    for k in ["skills_list", "languages", "soft_options"]:
        val= p.get(k,[])
        if isinstance(p.get(k), str):
            p[k] = [x.strip() for x in p[k].split(",") if x.strip()]
        elif isinstance(val , list):
            p[k] = val
        else:
            p[k] =[]
    # 🔧 FIX: Separate Technical vs Soft Skills
    tech_skills = []
    soft_skills = set(p.get("soft_options", []))  # preserve ATS soft skills

    for s in p.get("skills_list", []):
        skill = ""
        skill_type = ""

        if isinstance(s, dict):
            skill = s.get("skill", "").strip()
            skill_type = s.get("type", "").lower()
        else:
            skill = str(s).strip()
            skill_type = ""

        if not skill:
            continue

        if skill.lower() in SOFT_SKILL_KEYWORDS or skill_type == "soft":
            soft_skills.add(skill)
        else:
            tech_skills.append(skill)

    p["skills_list"] = sorted(set(tech_skills))
    p["soft_options"] = sorted(set(soft_skills))
    
    # 🔧 FIX LANGUAGES (ALL ATS FORMATS)
    langs = []

    raw_langs = p.get("languages", [])

# string → split
    if isinstance(raw_langs, str):
        raw_langs = [x.strip() for x in raw_langs.split(",") if x.strip()]

# list → normalize
    if isinstance(raw_langs, list):
        for l in raw_langs:
            if isinstance(l, dict):
                lang = l.get("language") or l.get("name") or ""
                if lang:
                    langs.append(lang.strip())
            elif isinstance(l, str):
                langs.append(l.strip())

    p["languages"] = sorted(set(langs))

    # Education
    edu_clean = []
    for e in p.get("education", []):
        if not isinstance(e, dict):
            continue
        edu_clean.append({
            "course": e.get("course",""),
            "school": e.get("school",""),
            "board": e.get("board",""),
            "startyear": e.get("startyear",""),
            "stopyear": e.get("stopyear",""),
            "sgpa": e.get("sgpa","")
        })
    p["education"] = edu_clean

    # Text fields
    for k in ["experience_raw","projects_raw","declaration_raw","summary"]:
        p[k] = p.get(k,"")

    return p

def setup_one_page(doc):
    section = doc.sections[0]

    # A4 Size
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # Tight margins
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def get_docx_bytes(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- Helper Function: Export to Docx ---
def create_docx(data):
    doc = Document()

    # THIS LINE MAKES RESUME 1 PAGE
    setup_one_page(doc)

    # ===== HEADER =====
    name_p = doc.add_heading(data["name"], level=0)
    name_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact_p = doc.add_paragraph(
        f"{data['email']} | {data['phone']} | {data['location']}"
    )
    contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ===== SUMMARY =====
    if data.get("summary"):
        h = doc.add_heading("Summary", level=1)
        for r in h.runs:
            r.font.size = Pt(11)

        p = doc.add_paragraph(data["summary"])
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1

    # ===== EDUCATION =====
    if data.get("education"):
        h = doc.add_heading("Education", level=1)
        for r in h.runs:
            r.font.size = Pt(11)

        for edu in data["education"]:
            if not any(v.strip() for v in edu.values()):
                continue

            years = ""
            if edu["startyear"] and edu["stopyear"]:
                years = f"({edu['startyear']} – {edu['stopyear']})"
            elif edu["startyear"] or edu["stopyear"]:
                years = f"({edu['startyear'] or edu['stopyear']})"

            p = doc.add_paragraph(
                f"{edu['course']} {years} | {edu['school']} | {edu['board']} | "
                f"SGPA: {edu['sgpa']}"
            )
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1

    # ===== TECHNICAL SKILLS =====
    h = doc.add_heading("Technical Skills", level=1)
    for r in h.runs:
        r.font.size = Pt(11)

    tech_ai = data.get("technical_skills_ai", "")
    if tech_ai:
        for line in tech_ai.split("•"):
            if line.strip():
                p = doc.add_paragraph(f"• {line.strip()}")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1

    # ===== EXPERIENCE =====
    if data.get("experience"):
        h = doc.add_heading("Experience", level=1)
        for r in h.runs:
            r.font.size = Pt(11)

        for line in data["experience"].split("•"):
            if line.strip():
                p = doc.add_paragraph(f"• {line.strip()}")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1

    # ===== PROJECTS =====
    if data.get("projects"):
        h = doc.add_heading("Projects", level=1)
        for r in h.runs:
            r.font.size = Pt(11)

        for line in data["projects"].split("•"):
            if line.strip():
                p = doc.add_paragraph(f"• {line.strip()}")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1

    # ===== DECLARATION =====
    if data.get("declaration"):
        h = doc.add_heading("Declaration", level=1)
        for r in h.runs:
            r.font.size = Pt(11)

        p = doc.add_paragraph(data["declaration"])
        p.paragraph_format.line_spacing = 1

    return doc

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def create_sidebar_docx(data):
    doc = Document()

    # FORCE 1 PAGE (MARGINS + A4)
    setup_one_page(doc)

    # ===== MAIN TABLE =====
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    left.width = Inches(2.4)
    right.width = Inches(4.6)

    # ===== LEFT SIDEBAR COLOR =====
    set_cell_bg(left, "2F3A40")

    # ===== HELPER FOR COMPACT TEXT =====
    def add_compact(cell, text, bold=False):
        p = cell.add_paragraph(text)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1
        if bold:
            p.runs[0].bold = True
        for r in p.runs:
            r.font.size = Pt(9)

    # ===== LEFT CONTENT =====
    name_p = left.paragraphs[0]
    run = name_p.add_run(data["name"].upper())
    run.bold = True
    run.font.size = Pt(12)

    add_compact(left, "\nCONTACT", bold=True)
    add_compact(left, data.get("location", ""))
    add_compact(left, str(data.get("phone", "")))
    add_compact(left, data.get("email", ""))

    add_compact(left, "\nSKILLS", bold=True)

    tech_ai = data.get("technical_skills_ai", "")
    if tech_ai:
        for line in tech_ai.split("•"):
            if line.strip():
                add_compact(left, f"• {line.strip()}")
    else:
        for skill in data.get("skills_list", [])[:8]:
            add_compact(left, f"• {skill}")

    if data.get("languages"):
        add_compact(left, "\nLANGUAGES", bold=True)
        for lang in data["languages"][:4]:
            add_compact(left, f"• {lang}")

    if data.get("soft_options"):
        add_compact(left, "\nSOFT SKILLS", bold=True)
        for skill in data["soft_options"][:6]:
            add_compact(left, f"• {skill}")

    # ===== RIGHT CONTENT =====
    def add_heading(cell, text):
        p = cell.add_paragraph(f"\n{text}")
        p.runs[0].bold = True
        for r in p.runs:
            r.font.size = Pt(11)

    def add_para(cell, text):
        p = cell.add_paragraph(text)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1
        for r in p.runs:
            r.font.size = Pt(9.5)

    if data.get("summary"):
        add_heading(right, "PROFESSIONAL SUMMARY")
        add_para(right, data["summary"])

    if data.get("experience"):
        add_heading(right, "EXPERIENCE")
        for line in data["experience"].split("•"):
            if line.strip():
                add_para(right, f"• {line.strip()}")

    if data.get("projects"):
        add_heading(right, "PROJECTS")
        for line in data["projects"].split("•"):
            if line.strip():
                add_para(right, f"• {line.strip()}")

    if data.get("education"):
        add_heading(right, "EDUCATION")
        for edu in data["education"]:
            if not any(v.strip() for v in edu.values()):
                continue

            years = ""
            if edu["startyear"] and edu["stopyear"]:
                years = f"({edu['startyear']} – {edu['stopyear']})"
            elif edu["startyear"] or edu["stopyear"]:
                years = f"({edu['startyear'] or edu['stopyear']})"

            edu_line = (
                f"{edu['course']} {years} | "
                f"{edu['school']} | {edu['board']} | "
                f"SGPA: {edu['sgpa']}"
            )
            add_para(right, edu_line)

    if data.get("declaration"):
        add_heading(right, "DECLARATION")
        add_para(right, data["declaration"])

    return doc

def create_modern_sidebar_docx(data):
    doc = Document()

    # FORCE SINGLE PAGE (A4 + TIGHT MARGINS)
    setup_one_page(doc)

    # ===== MAIN TABLE =====
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    left.width = Inches(4.6)
    right.width = Inches(2.4)

    # ===== RIGHT SIDEBAR COLOR =====
    set_cell_bg(right, "E9CBF2")

    # ===== COMPACT HELPERS =====
    def add_heading(cell, text):
        p = cell.add_paragraph(f"\n{text}")
        p.runs[0].bold = True
        for r in p.runs:
            r.font.size = Pt(11)

    def add_text(cell, text, size=9.5):
        p = cell.add_paragraph(text)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1
        for r in p.runs:
            r.font.size = Pt(size)

    # ===== LEFT MAIN CONTENT =====
    name_p = left.paragraphs[0]
    run = name_p.add_run(data["name"].upper())
    run.bold = True
    run.font.size = Pt(18)

    if data.get("summary"):
        add_heading(left, "PROFESSIONAL SUMMARY")
        add_text(left, data["summary"])

    if data.get("experience"):
        add_heading(left, "EXPERIENCE")
        for line in data["experience"].split("•"):
            if line.strip():
                add_text(left, f"• {line.strip()}")

    if data.get("projects"):
        add_heading(left, "PROJECTS")
        for line in data["projects"].split("•"):
            if line.strip():
                add_text(left, f"• {line.strip()}")

    if data.get("declaration"):
        add_heading(left, "DECLARATION")
        add_text(left, data["declaration"])

    # ===== RIGHT SIDEBAR CONTENT =====
    add_heading(right, "CONTACT")
    add_text(right, data.get("location", ""), size=9)
    add_text(right, str(data.get("phone", "")), size=9)
    add_text(right, data.get("email", ""), size=9)

    add_heading(right, "SKILLS")
    tech_ai = data.get("technical_skills_ai", "")
    if tech_ai:
        for line in tech_ai.split("•"):
            if line.strip():
                add_text(right, f"• {line.strip()}", size=9)
    else:
        for skill in data.get("skills_list", [])[:8]:
            add_text(right, f"• {skill}", size=9)

    if data.get("soft_options"):
        add_heading(right, "SOFT SKILLS")
        for skill in data["soft_options"][:6]:
            add_text(right, f"• {skill}", size=9)

    if data.get("languages"):
        add_heading(right, "LANGUAGES")
        for lang in data["languages"][:4]:
            add_text(right, f"• {lang}", size=9)

    if data.get("education"):
        add_heading(right, "EDUCATION")
        for edu in data["education"]:
            if not any(v.strip() for v in edu.values()):
                continue

            years = ""
            if edu["startyear"] and edu["stopyear"]:
                years = f"({edu['startyear']} – {edu['stopyear']})"
            elif edu["startyear"] or edu["stopyear"]:
                years = f"({edu['startyear'] or edu['stopyear']})"

            edu_line = (
                f"{edu['course']} {years} | "
                f"{edu['school']} | {edu['board']} | "
                f"SGPA: {edu['sgpa']}"
            )
            add_text(right, edu_line, size=9)

    return doc

LANGUAGE_OPTIONS = [
    # A
    "Afrikaans", "Akan", "Albanian", "Amharic", "Arabic", "Aragonese",
    "Armenian", "Assamese", "Aymara", "Azerbaijani",
    # B
    "Bambara", "Basque", "Belarusian", "Bengali", "Bhojpuri", "Bislama",
    "Bosnian", "Breton", "Bulgarian", "Burmese",
    # C
    "Catalan", "Cebuano", "Chamorro", "Chichewa",
    "Chinese (Mandarin)", "Chinese (Cantonese)", "Corsican",
    "Croatian", "Czech",
    # D
    "Danish", "Dhivehi", "Dogri", "Dutch", "Dzongkha",
    # E
    "English", "Esperanto", "Estonian", "Ewe",
    # F
    "Faroese", "Fijian", "Filipino", "Finnish", "French", "Frisian", "Fula",
    # G
    "Galician", "Georgian", "German", "Greek", "Greenlandic",
    "Guarani", "Gujarati",
    # H
    "Haitian Creole", "Hausa", "Hebrew", "Hindi", "Hmong", "Hungarian",
    # I
    "Icelandic", "Igbo", "Ilocano", "Indonesian", "Inuktitut",
    "Irish", "Italian",
    # J
    "Japanese", "Javanese",
    # K
    "Kannada", "Kazakh", "Khmer", "Kinyarwanda", "Korean",
    "Kurdish", "Kyrgyz",
    # L
    "Lao", "Latin", "Latvian", "Lingala", "Lithuanian", "Luxembourgish",
    # M
    "Macedonian", "Maithili", "Malagasy", "Malay", "Malayalam",
    "Maltese", "Maori", "Marathi", "Mongolian",
    # N
    "Nepali", "Newari", "Norwegian", "Nyanja",
    # O
    "Odia", "Oromo", "Ossetian",
    # P
    "Pashto", "Persian (Farsi)", "Polish", "Portuguese", "Punjabi",
    # Q
    "Quechua",
    # R
    "Romanian", "Russian",
    # S
    "Samoan", "Sanskrit", "Scots", "Scottish Gaelic", "Serbian",
    "Sesotho", "Setswana", "Shona", "Sindhi", "Sinhala",
    "Slovak", "Slovenian", "Somali", "Spanish", "Sundanese",
    "Swahili", "Swedish",
    # T
    "Tagalog", "Tajik", "Tamil", "Tatar", "Telugu", "Thai",
    "Tigrinya", "Tok Pisin", "Tongan", "Turkish", "Turkmen",
    # U
    "Ukrainian", "Urdu", "Uyghur", "Uzbek",
    # V
    "Vietnamese",
    # W
    "Welsh", "Wolof",
    # X
    "Xhosa",
    # Y
    "Yiddish", "Yoruba",
    # Z
    "Zulu"
]
SOFT_SKILL_OPTIONS = [
    "Communication","Teamwork","Leadership","Problem Solving","Critical Thinking","Time Management","Adaptability","Creativity",
    "Work Ethic","Attention to Detail","Decision Making","Conflict Resolution","Public Speaking","Emotional Intelligence",
    "Collaboration","Stress Management","Self Motivation","Active Listening","Negotiation","Flexibility"
]
st.set_page_config(page_title="Resume Builder", layout="centered")
#summary
def generate_summary_llama(data, user_summary=""):

    # 🔒 HARD BLOCK: template-style input (safety net)
    if user_summary and any(
        x in user_summary.lower()
        for x in ["[job title]", "[number", "[industry"]
    ):
        # In production, LOG instead of raising
        raise ValueError("Template-style output detected. Block generation.")

    raw_skills = data.get("skills_list", [])
    experience = data.get("experience", "")

    skills = ", ".join(
        s.get("skill", "") if isinstance(s, dict) else str(s)
        for s in raw_skills if s
    )

    # 🔹 CASE 1: Empty / Skip → FULL AUTO GENERATION
    if not user_summary or not user_summary.strip():
        prompt = f"""
Write a professional, ATS-friendly resume summary.

Rules:
- Do NOT add headings
- Do NOT use bullet points
- Avoid generic phrases
- Do NOT invent experience
- Return ONLY the summary text

Candidate Information:
Skills: {skills}
Experience: {experience}
"""
        return generate_ai_content(prompt)

    # 🔹 CASE 2: Short but intent-based ("i am brilliant")
    if is_intent_based_summary(user_summary):
        prompt = f"""
Professionally expand the following self-description into a resume summary.

STRICT RULES:
- Use ONLY the meaning of the user text
- Do NOT add years of experience
-do not add job title/ experience
-do not add area of skills 
-do not add [] type words
- Do NOT add skills, tools, or industries
- Do NOT add achievements or results
- Keep it neutral and fresher-safe
- ATS-friendly, plain sentences
- Return ONLY the summary text

User Text:
"{user_summary}"
"""
        return generate_ai_content(prompt)

    # 🔹 CASE 3: Very low quality junk
    if is_low_quality_summary(user_summary):
        prompt = f"""
Write a professional, ATS-friendly resume summary.

Rules:
- Neutral tone
- Fresher-safe
- No invented experience
-do not add years of experience
-do not add job title/ experience
-do not add area of skills 
-do not add [] type words
- Return ONLY the summary text
"""
        return generate_ai_content(prompt)

    # 🔹 CASE 4: Valid summary → Improve
    prompt = f"""
Rewrite and professionally improve the following resume summary.

Rules:
- Preserve original meaning
- Do NOT invent experience, skills, or achievements
- ATS-friendly
-do not add job title/ experience
-do not add area of skills 
-do not add [] type words
- Return ONLY the rewritten summary

User Summary:
"{user_summary}"
"""
    return generate_ai_content(prompt)
#def generate_unique_summary_from_input(user_summary: str) -> str:
def generate_resume_summary(user_input: str) -> str:
    if not user_input.strip():
        raise ValueError("Summary input is required")

    input_type = classify_input(user_input)
    style = random.choice(VARIATION_STYLES)

    prompt = f"""
Rewrite the following content into a PROFESSIONAL RESUME SUMMARY.

INPUT TYPE:
- {input_type}

STRICT RULES:
- Use ONLY information explicitly stated by the user
- Do NOT invent achievements, metrics, or responsibilities
- Do NOT add personality traits or motivation
- Keep tone professional and resume-appropriate
- Expand naturally to 3–4 lines
- ATS-safe wording
- Output must be unique on every generation
- Do NOT add job titles unless user mentions them

STYLE:
- {style}

User Content:
"{user_input}"

Return ONLY the resume summary.
"""

    response = ollama.chat(
        model="llama3.2:latest",
        messages=[
            {
                "role": "system",
                "content": "You generate factual resume summaries without assumptions."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return sanitize_summary(response["message"]["content"].strip())
#techincal skills
def generate_technical_llama(data):
    skills = ", ".join(data.get("skills_list", []))

    if not skills:
        return ""

    prompt = (
        "For each skill, add a 1-line professional description."
        "Rules:"
        "- Use bullet points only (•)"
        "- One line per skill"
        "- Return ONLY the bulleted list"
        f"Skills: {skills}"
    )

    return generate_ai_content(prompt)

#experience
def generate_experience_llama(data, is_fresher=False, years_of_exp=None, exp_text=""):
    if is_fresher:
        prompt = (
            "Generate 3 resume bullet points for a fresher based on internships,part-time jobs, or practical exposure."
            "Do NOT include academic or personal projects. "
            "Use bullet points only (•)"
            "No company names"
            "No years of experience"
            "ATS-friendly"
            "Return ONLY bullet points"
        )

    elif years_of_exp is not None and 0 < years_of_exp <= 3:
        prompt = (
            f"Generate 3 resume bullet points for a candidate with {years_of_exp} years of IT experience."
            "Use bullet points only (•)"
            "Focus on skills, tools, teamwork"
            "ATS-friendly"
            "Return ONLY bullet points"
        )

    elif exp_text.strip():
        prompt = (
            "Rewrite the following experience into 3 ATS-optimized resume bullet points."
            "Use bullet points only (•)"
            "Return ONLY bullet points"
            f"{exp_text}"
        )

    else:
        prompt = (
            "Generate exactly 3 resume bullet points based on technical skills and academic expsosure."
            "Rules:"
            "- Use bullet points only (•)"
            "- ATS-friendly"
            "- Return ONLY bullet points"
        )

    return generate_ai_content(prompt)

#projects
def generate_projects_llama(data, project_text=""):
    skills = ", ".join(data.get("skills_list", []))

    if project_text.strip():
        prompt = (
            "Rewrite the following into exactly 2 professional resume bullet points."
            "Rules:"
            "- Use bullet points only (•)"
            "- Focus on tools, technologies and impact"
            "- Do not mix with experience"
            "- Return ONLY bullet points"
            f"{project_text}"
        )
    else:
        prompt = (
            "Generate exactly 2 resume project bullet points."
            "Rules:"
            "- Use bullet points only (•)"
            "- ATS-friendly"
            "- Return ONLY bullet points"
            f"Skills: {skills}"
        )

    return generate_ai_content(prompt)

#declaration
def generate_declaration_llama(data, user_text=""):
    if user_text.strip():
        prompt = (
            "Rewrite the following resume declaration professionally."
            "1–2 lines only"
            "Formal tone"
            "Return ONLY the declaration text"
        )
    else:
        prompt = (
            "Write a professional resume declaration."
            "1-2 lines only"
            "formal tone"
            "Return only the declaration text."
            )

    return generate_ai_content(prompt)


# ---------- SESSION STATE ----------
if "page" not in st.session_state:
    st.session_state.page = "home"

if "resume_type" not in st.session_state:
    st.session_state.resume_type = None

if "template" not in st.session_state:
    st.session_state.template = None

# ---------- FORM SESSION STATE ----------
if "form_step" not in st.session_state:
    st.session_state.form_step = 1

if "form_data" not in st.session_state:
    st.session_state.form_data = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "",
        "education": [],
        "skills_list": [],
        "languages": [],
        "soft_options": [],
        "experience_raw": "",
        "projects_raw": "",
        "declaration_raw": ""
    }
defaults = {
    "name_input": "",
    "email_input": "",
    "phone_input": "",
    "location_input": "",
    "summary_input": "",
    "experience_input": "",
    "projects_input": "",
    "declaration_input": "",
    "skills_input": [],
    "languages_input": [],
    "soft_input": []
}

for k,v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v
# ---- WIDGET STATE FOR AUTO-FILL ----
for key in ["name", "email", "phone", "location"]:
    widget_key = f"{key}_input"
    if widget_key not in st.session_state:
        st.session_state[widget_key] = ""

if "uploaded_resume" not in st.session_state:
    st.session_state.uploaded_resume = None

if "education_rows" not in st.session_state:
    st.session_state.education_rows = 1


# ---------- HOME PAGE ----------
if st.session_state.page == "home":

    st.markdown(
        "<h1 style='text-align:center; color:#1f4fd8;'>Resume Builder</h1>",
        unsafe_allow_html=True
    )

    col1, col2 = st.columns(2)

    with col1:
        if st.button("New Resume", use_container_width=True):
            st.session_state.resume_type = "new"

    with col2:
        if st.button("Existing Resume", use_container_width=True):
            st.session_state.resume_type = "existing"

    st.markdown("<br>", unsafe_allow_html=True)

    if st.session_state.resume_type:
        if st.button("--> Next", use_container_width=True):
            if st.session_state.resume_type == "new":
                st.session_state.page = "templates"

            elif st.session_state.resume_type == "existing":
                st.session_state.page = "upload"

            st.rerun()

# ---------- UPLOAD EXISTING RESUME ----------
elif st.session_state.page == "upload":

    st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Upload Existing Resume</h2>",
        unsafe_allow_html=True
    )

    uploaded_file = st.file_uploader(
        "Upload your resume (PDF or DOCX)",
        type=["pdf", "docx"]
    )

    st.markdown("<br>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("<-- Back"):
            st.session_state.page = "home"
            st.rerun()

    if uploaded_file:
        if st.button("--> Continue", use_container_width=True):

            with st.spinner("Running ATS analysis..."):

                # 1️⃣ Extract text
                resume_text = extract_resume_text(uploaded_file)
                if len(resume_text.strip()) < 200:
                    st.warning(
        "This resume appears to be scanned or image-based. "
        "Autofill may be limited. Please review manually."
                    )
                ats_output = ats_parse_resume(resume_text)
                # 2️⃣ ATS parse
                # 1️⃣ ATS PARSE + NORMALIZE
                parsed = normalize_ats_data(ats_output)
                
                parsed.setdefault("experience_raw", "")
                parsed.setdefault("projects_raw", "")
                parsed.setdefault("declaration_raw", "")
                parsed.setdefault("summary", "")

                contact = extract_contact_regex(resume_text)
                parsed["email"] = parsed.get("email") or contact.get("email", "")
                parsed["phone"] = parsed.get("phone") or contact.get("phone", "")
                parsed["name"]  = parsed.get("name")  or contact.get("name", "")
                safe_location = extract_location_safely(resume_text)
                parsed["location"] = safe_location or parsed.get("location", "")


# 2️⃣ UPDATE MAIN FORM DATA (THIS FEEDS ALL STEPS)
                st.session_state.form_data.update(parsed)

# 3️⃣ PERSONAL DETAILS → WIDGET STATE
                st.session_state.name_input = parsed.get("name", "")
                st.session_state.email_input = parsed.get("email", "")
                st.session_state.phone_input = str(parsed.get("phone", ""))
                st.session_state.location_input = parsed.get("location", "")

# 4️⃣ OTHER SECTIONS → WIDGET STATE
                st.session_state.skills_input = parsed.get("skills_list", [])
                st.session_state.languages_input = parsed.get("languages", [])
                st.session_state.soft_input = parsed.get("soft_options", [])

# 5️⃣ EDUCATION → AUTO ROWS
                st.session_state.education_rows = max(1, len(parsed.get("education", [])))

# 6️⃣ RAW TEXT FIELDS
                st.session_state.summary_input = parsed.get("summary", "")
                st.session_state.experience_input = parsed.get("experience_raw", "")
                st.session_state.projects_input = parsed.get("projects_raw", "")
                st.session_state.declaration_input = parsed.get("declaration_raw", "")

# 7️⃣ MOVE TO FORM PAGE
                st.session_state.page = "form"
                st.session_state.form_step = 1
                st.rerun()
                    #  Move to template selection
                    #st.session_state.page = "form"
                    #st.rerun()
# ---------- TEMPLATE SELECTION PAGE ----------
elif st.session_state.page == "templates":

    st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Choose Resume Template</h2>",
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        st.image("res_template1.png")
        if st.button("Simple", use_container_width=True):
            st.session_state.template = "simple"
            st.session_state.page = "form"
            st.session_state.form_step = 1
            st.rerun()

        
    with col2:
        st.image("res_template2.png")
        if st.button("Modern", use_container_width=True):
            st.session_state.template = "modern"
            st.session_state.page = "form"
            st.session_state.form_step = 1
            st.rerun()

    with col3:
        st.image("res_template3.png")
        if st.button("Sidebar", use_container_width=True):
            st.session_state.template = "sidebar"
            st.session_state.page = "form"
            st.session_state.form_step = 1
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("<-- Back"):
            st.session_state.page = "home"
            st.rerun()

    with col2:
        if st.session_state.template:
            if st.button("--> Next", use_container_width=True):
                st.session_state.page = "form"
                st.rerun()

elif st.session_state.page == "form":

    # ---------- STEP 1 : PERSONAL DETAILS ----------
    if st.session_state.form_step == 1:

        st.markdown(
            "<h2 style='text-align:center; color:#1f4fd8;'>Personal Details</h2>",
            unsafe_allow_html=True
        )

        st.subheader("Personal Details")
        # SAFELY INIT WIDGET STATE FROM ATS (ONLY ONCE)
        for field in ["name", "email", "phone", "location"]:
            widget_key = f"{field}_input"
            if widget_key not in st.session_state or not st.session_state[widget_key]:
                st.session_state[widget_key] = st.session_state.form_data.get(field, "")

        st.text_input("Full Name", key="name_input")
        st.text_input("Email", key="email_input")
        st.text_input("Phone Number", key="phone_input", max_chars=10)
        st.text_input("Location", key="location_input")

        #  Sync widgets → form_data
        st.session_state.form_data.update({
            "name": st.session_state.name_input,
            "email": st.session_state.email_input,
            "phone": st.session_state.phone_input,
            "location": st.session_state.location_input
        })

        phone_valid = False
        if st.session_state.phone_input:
            if st.session_state.phone_input.isdigit() and len(st.session_state.phone_input) == 10:
                phone_valid = True
                st.session_state.form_data["phone"] = int(st.session_state.phone_input)
                st.success("Valid phone number")
            else:
                st.error("Phone number must be exactly 10 digits")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("<-- Back"):
                st.session_state.page = "templates"
                st.session_state.form_step = 1
                st.rerun()

        with col2:
            if st.button("--> Next"):

                if not st.session_state.form_data["name"].strip():
                    st.error("Enter your name. It is compulsory")
                    st.stop()

                if not st.session_state.form_data["email"].strip():
                    st.error("Enter your email. It is compulsory")
                    st.stop()

                if not st.session_state.phone_input.strip():
                    st.error("Enter your phone number")
                    st.stop()

                if not (st.session_state.phone_input.isdigit() and len(st.session_state.phone_input) == 10):
                    st.error("Phone number must be exactly 10 digits")
                    st.stop()

                st.session_state.form_data["phone"] = st.session_state.phone_input
                st.session_state.form_step = 2
                st.rerun()

    if st.session_state.form_step == 2:
        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Professional Summary</h2>",
        unsafe_allow_html=True)
        if "summary_input" not in st.session_state:
            st.session_state.summary_input = st.session_state.form_data.get("summary", "")

        st.subheader("Summary (Optional)")

        st.text_area(
    "Write your summary",
    height=150,
    key="summary_input"
)

        st.session_state.form_data["summary"] = st.session_state.summary_input

        col1, col2, col3 = st.columns(3)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 1
                st.rerun()

    # ---------- SKIP ----------
        with col2:
            if st.button(">> Skip"):
                st.session_state.form_data["summary"] = generate_summary_llama(
                    st.session_state.form_data
                )
                st.session_state.form_step = 3
                st.rerun()

    # ---------- NEXT ----------
        with col3:
            if st.button("--> Next"):
                user_summary = st.session_state.summary_input.strip()

                if user_summary:
            # 🔑 USER-BASED UNIQUE REWRITE
                    st.session_state.form_data["summary"] = generate_resume_summary(
                        user_summary
                    )
                else:
            #  FALLBACK: auto-generate if empty
                    st.session_state.form_data["summary"] = generate_summary_llama(
                        st.session_state.form_data
                    )

                st.session_state.form_step = 3
                st.rerun()

    elif st.session_state.form_step == 3:
        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Education</h2>",
        unsafe_allow_html=True)

        st.subheader("Education (Optional)")

        education_data = []

        education_data = []

        for i in range(st.session_state.education_rows):
            st.markdown(f"**Education {i+1}**")

            edu = st.session_state.form_data.get("education", [])
            edu_data = edu[i] if i < len(edu) else {}

            course = st.text_input("Course / Degree", value=edu_data.get("course", ""), key=f"course_{i}")
            school = st.text_input("School / College", value=edu_data.get("school", ""), key=f"school_{i}")
            board = st.text_input("Board / University", value=edu_data.get("board", ""), key=f"board_{i}")
            startyear = st.text_input("Start Year", value=edu_data.get("startyear", ""), max_chars=4,key=f"start_{i}")
            stopyear = st.text_input("End Year", value=edu_data.get("stopyear", ""), max_chars=4,key=f"stop_{i}")
            sgpa = st.text_input("SGPA / Percentage", value=edu_data.get("sgpa", ""), max_chars=5,key=f"sgpa_{i}")

            education_data.append({
        "course": course,
        "school": school,
        "board": board,
        "startyear": startyear,
        "stopyear": stopyear,
        "sgpa": sgpa
    })


        col_back, col_add, col_skip, col_next = st.columns([1, 2, 1, 1])

    #  Add more education
        with col_add:
            if st.button("Add More"):
                st.session_state.education_rows += 1
                st.rerun()

    # ⬅ Back
        with col_back:
            if st.button("<-- Back"):
                st.session_state.form_step = 2
                st.rerun()

    # ⏭ Skip
        with col_skip:
            if st.button(">> Skip"):
                st.session_state.form_data["education"] = []
                st.session_state.form_step = 4
                st.rerun()

    # ➡ Next
        with col_next:
            if st.button("--> Next"):

                clean_education = []
                for edu in education_data:
                    if any(v.strip() for v in edu.values()):
                        clean_education.append(edu)

                for edu in clean_education:
                    start = edu.get("startyear", "").strip()
                    stop = edu.get("stopyear", "").strip()
                    course = edu.get("course", "Education")

                    if start and (not start.isdigit() or len(start) != 4):
                        st.error(f"Invalid start year for {course}")
                        st.stop()

                    if stop and (not stop.isdigit() or len(stop) != 4):
                        st.error(f"Invalid stop year for {course}")
                        st.stop()

                    if start and stop:
                        s, e = int(start), int(stop)
                        if e < s:
                            st.error(f"End year cannot be before start year for {course}")
                            st.stop()

                st.session_state.form_data["education"] = clean_education
                st.session_state.form_step = 4
                st.rerun()

    
    # ---------- STEP 4 : TECHNICAL SKILLS (MANDATORY) ----------
    elif st.session_state.form_step == 4:
        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Technical Skills</h2>",
        unsafe_allow_html=True
    )

        st.subheader("Technical Skills")

        skills_input = st.text_area(
        "Enter at least ONE technical skill (comma separated)",
        placeholder="Python, SQL, Streamlit, Git",
        height=120,
        value=", ".join(st.session_state.form_data.get("skills_list", []))
    )

    # ---------- VALIDATION ----------
        skills_list = []
        skills_valid = False

        if skills_input.strip():
            skills_list = [s.strip() for s in skills_input.split(",") if s.strip()]
            if len(skills_list) >= 1:
                skills_valid = True
                st.success(f" {len(skills_list)} skill(s) added")
            else:
                st.error("Please enter at least one valid skill")
        else:
            st.warning("At least one technical skill is required")

        col1, col2 = st.columns(2)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 3
                st.rerun()

    # ---------- NEXT ----------
       # ---------- NEXT ----------
        with col2:
            if st.button("--> Next", disabled=not skills_valid):
                st.session_state.form_data["skills_list"] = skills_list

                st.session_state.form_data["technical_skills_ai"] = generate_technical_llama(
                    st.session_state.form_data
                )

                st.session_state.form_step = 5
                st.rerun()

    elif st.session_state.form_step == 5:
        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Languages</h2>",
        unsafe_allow_html=True
    )

        st.subheader("Languages Known (Optional)")

        st.multiselect(
        "Select all languages you know",
        options=LANGUAGE_OPTIONS,
        key="languages_input"
    )

    #  Sync widget → form_data
        st.session_state.form_data["languages"] = st.session_state.languages_input
        col1, col2 = st.columns(2)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 4
                st.rerun()

    # ---------- NEXT ----------
        with col2:
            if st.button("--> Next"):
                #st.session_state.form_data["languages"] = selected_languages
                st.session_state.form_step = 6
                st.rerun()
     
    # ---------- STEP 6 : SOFT SKILLS (OPTIONAL) ----------
    elif st.session_state.form_step == 6:

        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Soft Skills</h2>",
        unsafe_allow_html=True)

        st.subheader("Soft Skills")

        st.multiselect(
    "Select your soft skills",
    options=SOFT_SKILL_OPTIONS,
    key="soft_input"
)

        st.session_state.form_data["soft_options"] = st.session_state.soft_input
        col1, col2, col3 = st.columns(3)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 5
                st.rerun()

    # ---------- SKIP ----------
        with col2:
            if st.button(">> Skip"):
                st.session_state.form_data["soft_options"] = []
                st.session_state.form_step = 7
                st.rerun()

    # ---------- NEXT ----------
        with col3:
            if st.button("--> Next"):
                #st.session_state.form_data["soft_options"] = selected_soft_skills
                st.session_state.form_step = 7
                st.rerun()

    # ---------- STEP 7 : EXPERIENCE (OPTIONAL & SMART) ----------
    elif st.session_state.form_step == 7:
        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Experience</h2>",
        unsafe_allow_html=True
    )

        st.subheader("Professional Experience")

    # ---------- EXPERIENCE LEVEL ----------
        experience_level = st.selectbox(
            "Experience Level",
        [
            "Fresher",
            "1 Year",
            "2 Years",
            "3 Years",
            "5+ Years",
            "10+ Years"
        ],
            index=0
    )
        
    # ---------- EXPERIENCE INPUT ----------
        experience_input = st.text_area(
        "Describe your experience (optional)",
            placeholder=(
            "Example:\n"
            "- Worked as Python Developer at XYZ\n"
            "- Built REST APIs\n"
            "- Collaborated with team\n\n"
            "Leave empty to auto-generate"
            ),
            height=180,
            value=st.session_state.form_data.get("experience_raw", "")
        )

        col1, col2, col3 = st.columns(3)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 6
                st.rerun()

    # ---------- SKIP ----------
        with col2:
            if st.button(">> Skip"):
                st.session_state.form_data["experience"] = generate_experience_llama(
                    st.session_state.form_data,
                    is_fresher=(experience_level == "Fresher")
            )
                st.session_state.form_step = 8
                st.rerun()

    # ---------- NEXT ----------
        with col3:
            if st.button("--> Next"):

                exp_text = experience_input.strip()

                years_of_exp = None
                if exp_text.isdigit():
                    years_of_exp = int(exp_text)

                st.session_state.form_data["experience"] = generate_experience_llama(
                    st.session_state.form_data,
                    is_fresher=(experience_level == "Fresher"),
                    years_of_exp=years_of_exp,
                    exp_text=exp_text
                    )

                st.session_state.form_data["experience_raw"] = exp_text

                st.session_state.form_step = 8
                st.rerun()

    # ---------- STEP 8 : PROJECTS (OPTIONAL) ----------
    elif st.session_state.form_step == 8:

        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Projects</h2>",
        unsafe_allow_html=True
    )

        st.subheader("Projects (Optional)")

        projects_input = st.text_area(
        "Describe your projects (optional)",
            placeholder=(
            "Example:\n"
            "- AI Resume Builder using Python & Streamlit\n"
            "- E-commerce Website using Django\n\n"
            "Leave empty to skip"
        ),
            height=200,
            value=st.session_state.form_data.get("projects_raw", "")
    )

        col1, col2, col3 = st.columns(3)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 7
                st.rerun()

    # ---------- SKIP ----------
        with col2:
            if st.button(">> Skip"):
                st.session_state.form_data["projects"] = ""   # IMPORTANT
                st.session_state.form_step = 9
                st.rerun()

    # ---------- NEXT ----------
        with col3:
            if st.button("--> Next"):
                if projects_input.strip():
                    st.session_state.form_data["projects"] = generate_projects_llama(
                        st.session_state.form_data,
                        projects_input
                    )
                else:
                    st.session_state.form_data["projects"] = ""
                st.session_state.form_step = 9
                st.session_state.form_data["projects_raw"] = projects_input
                st.rerun()

    # ---------- STEP 9 : DECLARATION ----------
    elif st.session_state.form_step == 9:

        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Declaration</h2>",
        unsafe_allow_html=True
    )

        st.subheader("Declaration (Optional)")

        declaration_input = st.text_area(
        "Write your declaration (optional)",
            placeholder=(
            "Example:\n"
            "I hereby declare that the above information is true and correct to "
            "the best of my knowledge and belief."
        ),
            height=120,
            value=st.session_state.form_data.get("declaration_raw", "")
        )

        col1, col2, col3 = st.columns(3)

    # ---------- BACK ----------
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 8
                st.rerun()

    # ---------- SKIP ----------
        with col2:
            if st.button(">> Skip"):
                st.session_state.form_data["declaration"] = generate_declaration_llama(
                    st.session_state.form_data
            )
                st.session_state.form_step = 10
                st.rerun()

    # ---------- NEXT ----------
        with col3:
            if st.button("--> Next"):
                if declaration_input.strip():
                    st.session_state.form_data["declaration"] = declaration_input.strip()
                else:
                    st.session_state.form_data["declaration"] = generate_declaration_llama(
                        st.session_state.form_data
                    )
                st.session_state.form_step = 10
                st.rerun()

    # ---------- STEP 10 : PREVIEW ----------
    elif st.session_state.form_step == 10:

        st.markdown(
        "<h2 style='text-align:center; color:#1f4fd8;'>Resume Preview</h2>",
            unsafe_allow_html=True
        )

        data = st.session_state.form_data
        required_fields = ["name", "email", "phone", "skills_list"]
        for field in required_fields:
            if not data.get(field):
                st.error(f"Missing required field: {field}")
                st.stop()

    # ---------- BASIC PREVIEW ----------
        st.subheader("Personal Information")
        st.write(f"**Name:** {data.get('name','')}")
        st.write(f"**Email:** {data.get('email','')}")
        st.write(f"**Phone:** {data.get('phone','')}")
        st.write(f"**Location:** {data.get('location','')}")

        st.markdown("---")

    # ---------- SUMMARY ----------
        if data.get("summary"):
            st.subheader("Summary")
            st.write(data["summary"])

    # ---------- EDUCATION ----------
        if data.get("education"):
            st.subheader("Education")
            for edu in data["education"]:
                if not any(v.strip() for v in edu.values()):
                    continue
                st.write(
                    f"**{edu.get('course','')}** "
                    f"({edu.get('startyear','')} – {edu.get('stopyear','')})"
                )
                st.write(
                f"{edu.get('school','')} | {edu.get('board','')}"
                )
                if edu.get("sgpa"):
                    st.write(f"SGPA/Percentage: {edu.get('sgpa')}")
                    st.markdown("")

    # ---------- TECHNICAL SKILLS ----------
        st.subheader("Technical Skills")
        tech_ai = data.get("technical_skills_ai", "")
        if tech_ai:
            for line in tech_ai.split("•"):
                if line.strip():
                    st.write(f"• {line.strip()}")
        else:
            st.write(", ".join(data.get("skills_list", [])))

    # ---------- LANGUAGES ----------
        if data.get("languages"):
            st.subheader("Languages")
            st.write(", ".join(data["languages"]))

    # ---------- SOFT SKILLS ----------
        if data.get("soft_options"):
            st.subheader("Soft Skills")
            st.write(", ".join(data["soft_options"]))

    # ---------- EXPERIENCE ----------
        if data.get("experience"):
            st.subheader("Experience")
            st.write(data["experience"])

    # ---------- PROJECTS ----------
        if data.get("projects"):
            st.subheader("Projects")
            st.write(data["projects"])

    # ---------- DECLARATION ----------
        if data.get("declaration"):
            st.subheader("Declaration")
            st.write(data["declaration"])

        st.markdown("---")

    # ---------- ACTION BUTTONS ----------
        col1, col2 = st.columns(2)

    # ⬅ BACK
        with col1:
            if st.button("<-- Back"):
                st.session_state.form_step = 9
                st.rerun()

    # ⬇ DOWNLOAD
        with col2:
            template = st.session_state.template
            if template == "simple":
                doc = create_docx(data)
            elif template == "sidebar":
                doc = create_sidebar_docx(data)
            elif template == "modern":
                doc = create_modern_sidebar_docx(data)
            else:
                doc = create_docx(data)  

            doc_bytes = get_docx_bytes(doc)

            st.download_button(
                label="Download Resume (DOCX)",
                data=doc_bytes,
                file_name="resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )



